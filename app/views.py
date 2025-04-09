# from django.shortcuts import render, redirect
# from .models import Question, UploadBatch
# from .forms import UploadDocForm
# import docx
# import openai
# from django.conf import settings
# from django.contrib import messages
# from docx import Document

# # Set API key directly
# openai.api_key = settings.OPENAI_API_KEY

# def home(request):
#     questions = Question.objects.all()
#     return render(request, 'home.html', {'questions': questions})

# def upload_docx(request):
#     if request.method == 'POST':
#         form = UploadDocForm(request.POST, request.FILES)
#         if form.is_valid():
#             doc_file = request.FILES['file']
#             doc = docx.Document(doc_file)
#             for para in doc.paragraphs:
#                 if para.text.strip():
#                     Question.objects.get_or_create(text=para.text.strip())
#             return redirect('home')
#     else:
#         form = UploadDocForm()
#     return render(request, 'upload.html', {'form': form})


# def generate_answers(request):
#     questions = Question.objects.filter(is_answered=False)
#     count = questions.count()
#     for q in questions:
#         prompt = f"Answer the following question in detail:\n{q.text}"
#         try:
#             response = openai.ChatCompletion.create(
#                 model="gpt-3.5-turbo",
#                 messages=[
#                     {"role": "system", "content": "You are a helpful assistant."},
#                     {"role": "user", "content": prompt},
#                 ]
#             )
#             answer = response['choices'][0]['message']['content'].strip()
#             # answer = "answer"+str(count)
            
#             # count -=1
#             q.answer = answer
#             q.is_answered = True

#             q.save()
#         except Exception as e:
#             print(f"Error for question: {q.text[:30]}... => {e}")
#     return redirect('home')



### views.py

from django.shortcuts import render, redirect
from django.contrib import messages
from .models import Question, UploadBatch
from docx import Document
import openai
import os

# Set your OpenAI API key
openai.api_key = os.getenv("OPENAI_API_KEY")  # or use a hardcoded key for testing


def get_openai_answer(question_text):
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a helpful assistant that provides short, clear answers."},
                {"role": "user", "content": question_text},
            ]
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Error: {e}"


def upload_questions(request):
    if request.method == "POST" and request.FILES.get("docx_file"):
        doc_file = request.FILES["docx_file"]
        batch = UploadBatch.objects.create(filename=doc_file.name)

        document = Document(doc_file)
        added_count = 0

        for para in document.paragraphs:
            text = para.text.strip()
            
            if text and not Question.objects.filter(batch=batch, text=text).exists():
                # answer = get_openai_answer(text)
                answer = "answer"+str(added_count)  # Placeholder for OpenAI answer
                Question.objects.create(batch=batch, text=text, answer=answer)
                added_count += 1

        messages.success(request, f"{added_count} new questions with answers added under batch {batch.id}.")
        return redirect("view_batch", batch_id=batch.id)

    return render(request, "upload.html")


def view_batch(request, batch_id):
    batch = UploadBatch.objects.get(id=batch_id)
    questions = batch.questions.all()
    return render(request, "view_batch.html", {"batch": batch, "questions": questions})
