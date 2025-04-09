from docx import Document

# Create a new Word document
doc = Document()
doc.add_heading("Tech Interview Questions", level=1)

questions = [
    "What is a primary key in a relational database?",
    "Explain the difference between SQL and NoSQL databases.",
    "What are ACID properties in databases?",
    
]

# Add questions to the document
for i, q in enumerate(questions, 1):
    doc.add_paragraph(f"{i}. {q}")

# Save the document
doc_path = "tech_questions_50.docx"
doc.save(doc_path)


